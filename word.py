from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import re # For parsing table data

# --- Partnership Agreement Text ---
# (Same text as before)
agreement_text = """
GENERAL PARTNERSHIP AGREEMENT
OF
DRIP-O-MATIC

This General Partnership Agreement (the "Agreement") is made and entered into this ______ day of _______________, 2024 (the "Effective Date"), by and among the following individuals (collectively referred to as the "Partners," and individually as a "Partner"):

1.  Stella Marie B. Abanono
2.  Nathaniel G. Bingco
3.  Samantha Jane S. Darullo
4.  Louize Christopher S. Galang
5.  Yuri B. Gatbunton
6.  Neil Vincent E. Jaen
7.  Myke Jonard S. Sentillas

RECITALS

WHEREAS, the Partners desire to associate themselves as partners in a general partnership for the purpose of engaging in the business described herein; and

WHEREAS, the Partners wish to establish their respective rights, duties, and liabilities relating to the Partnership and its business;

NOW, THEREFORE, in consideration of the mutual covenants contained herein, the Partners agree as follows:

ARTICLE I: FORMATION, NAME, AND PURPOSE

1.1 Formation: The Partners hereby form a general partnership (the "Partnership") pursuant to the laws of the Republic of the Philippines.
1.2 Name: The name of the Partnership shall be Drip-o-matic. The business of the Partnership shall be conducted under this name or such other names as the Partners may unanimously agree upon.
1.3 Purpose: The primary purpose of the Partnership is to design, develop, manufacture, market, and sell the "Drip-o-matic" automatic plant care system and related products and services, including automatic irrigation, soil moisture monitoring, temperature regulation, plant misting, and solar power integration, aiming to provide convenient and efficient plant care solutions. The Partnership may also engage in any other lawful business activities related to its primary purpose as the Partners may unanimously agree.
1.4 Principal Place of Business: The principal place of business of the Partnership shall initially be located at STI Academic Center Altaraza Town Center, Tungkong Mangga, CSJDM, or such other location(s) as the Partners may unanimously determine.

ARTICLE II: TERM

2.1 Duration: The Partnership shall commence on the Effective Date and shall continue until dissolved and terminated in accordance with the provisions of this Agreement or by operation of law.

ARTICLE III: CAPITAL CONTRIBUTIONS

3.1 Initial Contributions: Each Partner shall contribute capital to the Partnership as mutually agreed upon. The description and agreed value of the initial capital contributions (which may include cash, property, services, or expertise) made by each Partner shall be detailed in Exhibit A, attached hereto and incorporated herein by reference.
3.2 Additional Contributions: No Partner shall be required to make additional capital contributions unless unanimously agreed upon in writing by all Partners.
3.3 No Interest: No Partner shall be entitled to receive interest on their capital contributions.
3.4 Capital Accounts: An individual capital account shall be maintained for each Partner, reflecting their contributions, share of profits and losses, and distributions.

ARTICLE IV: PROFITS, LOSSES, AND DISTRIBUTIONS

4.1 Allocation of Profits and Losses: Unless otherwise unanimously agreed in writing, the net profits and losses of the Partnership (after deduction of all expenses, including any salaries approved under Article VI) shall be allocated among the Partners in equal shares (i.e., 1/7th each).
4.2 Distributions: Distributions of cash or other property shall be made to the Partners at such times and in such amounts as the Partners may unanimously determine, typically reviewed on a quarterly basis, provided that such distributions do not impair the ability of the Partnership to meet its obligations. Distributions shall be made in proportion to the Partners' respective shares of profits as defined in Section 4.1.

ARTICLE V: MANAGEMENT AND DUTIES

5.1 Management: All Partners shall have the right to participate in the management and control of the Partnership's business.
5.2 Officer Roles: The Partners agree to designate initial officer roles to manage specific operational areas, subject to the overall authority of the Partnership:
    *   Chief Executive Officer (CEO): Yuri B. Gatbunton
    *   Chief Operating Officer (COO): Neil Vincent E. Jaen
    *   Chief Financial Officer (CFO): Samantha Jane S. Darullo
    *   Chief Marketing Officer (CMO): Stella Marie B. Abanono
    *   Chief Technology Officer (CTO): Louize Christopher S. Galang
    *   Chief Sales Officer (CSO): Myke Jonard S. Sentillas
    *   Chief Human Resources Officer (CHRO): Nathaniel G. Bingco
    The specific duties associated with these roles shall be as generally understood for such positions and as further defined by unanimous agreement of the Partners. These are internal management roles; all Partners remain general partners with associated rights and liabilities.
5.3 Decision Making:
    (a) Ordinary Course Decisions: Decisions arising in the ordinary course of the Partnership's business may be made by the designated officer responsible for that area, subject to review by the Partners.
    (b) Major Decisions: The following decisions shall require the unanimous written consent of all Partners:
        (i) Admitting a new Partner;
        (ii) Amending this Agreement;
        (iii) Selling, leasing, or disposing of all or substantially all of the Partnership's assets;
        (iv) Incurring debt or making expenditures outside the approved budget or above a threshold amount of [Specify Amount, e.g., PHP 50,000];
        (v) Changing the primary purpose or nature of the Partnership's business;
        (vi) Merging or consolidating the Partnership with another entity;
        (vii) Filing for bankruptcy or initiating dissolution;
        (viii) Approving annual budgets and major strategic plans.
5.4 Devotion of Time: Each Partner shall devote such time and attention to the Partnership business as is reasonably necessary for the discharge of their agreed-upon duties and responsibilities.

ARTICLE VI: SALARIES AND COMPENSATION

6.1 Salaries: Initially, no Partner shall receive a fixed salary for services rendered to the Partnership. Partners shall be compensated through their share of profits as provided in Article IV. The Partners may unanimously agree in writing to provide salaries or guaranteed payments to one or more Partners at a future date.
6.2 Expense Reimbursement: Partners shall be reimbursed for reasonable and necessary out-of-pocket expenses incurred directly on behalf of the Partnership business, provided such expenses are properly documented and approved according to procedures established by the Partners.

ARTICLE VII: BOOKS, RECORDS, AND BANKING

7.1 Books and Records: The Partnership shall maintain complete and accurate books and records of its business and financial affairs at its principal place of business, in accordance with generally accepted accounting principles (GAAP). Each Partner shall have the right to inspect and copy these records at reasonable times.
7.2 Fiscal Year: The fiscal year of the Partnership shall end on December 31st.
7.3 Banking: All funds of the Partnership shall be deposited in the Partnership's name in such bank account(s) as selected by the Partners. Withdrawals from such accounts shall be made upon the signature(s) of such Partner(s) as designated by unanimous agreement (initially proposed: CEO and CFO jointly, or as otherwise agreed).

ARTICLE VIII: PARTNER WITHDRAWAL, RETIREMENT, DEATH

8.1 Voluntary Withdrawal: Any Partner may withdraw from the Partnership upon providing at least ninety (90) days' prior written notice to the other Partners.
8.2 Buy-Out Option: Upon the withdrawal, retirement, death, or incapacity of a Partner (the "Departing Partner"), the remaining Partners shall have the option, exercisable within sixty (60) days of the event, to purchase the Departing Partner's interest in the Partnership.
8.3 Valuation: The purchase price for the Departing Partner's interest shall be the value of their capital account as of the date of departure, adjusted for profits/losses up to that date, plus or minus any other adjustments agreed upon by the Departing Partner (or their estate) and the remaining Partners. If an agreement on value cannot be reached, the value shall be determined by an independent appraiser mutually acceptable to both parties.
8.4 Payment Terms: The purchase price shall be paid by the remaining Partners (or the Partnership, if agreed) over a period not to exceed [Specify Period, e.g., twelve (12) months], with terms to be mutually agreed upon.
8.5 Dissolution upon Departure: If the remaining Partners do not exercise the option to purchase the Departing Partner's interest, the Partnership may be dissolved in accordance with Article IX.

ARTICLE IX: DISSOLUTION AND LIQUIDATION

9.1 Events of Dissolution: The Partnership shall be dissolved upon the occurrence of any of the following:
    (a) The unanimous written agreement of all Partners to dissolve;
    (b) The expiration of the term, if specified and not extended;
    (c) The sale or disposition of all or substantially all Partnership assets;
    (d) An event making it unlawful for the business of the Partnership to be carried on;
    (e) Entry of a decree of judicial dissolution;
    (f) Failure of remaining partners to purchase a Departing Partner's interest as per Section 8.5, if they elect dissolution.
9.2 Liquidation: Upon dissolution, the Partnership shall cease carrying on business except as necessary for winding up its affairs. The Partners (or a designated liquidating Partner/trustee) shall proceed with reasonable promptness to liquidate the Partnership's assets.
9.3 Distribution of Assets: The proceeds from liquidation shall be applied in the following order of priority:
    (a) To pay debts and liabilities of the Partnership owed to third-party creditors;
    (b) To pay debts and liabilities owed to Partners other than for capital and profits;
    (c) To repay Partners' capital contributions (adjusted for profits/losses);
    (d) Any remaining amount shall be distributed to the Partners according to their profit-sharing ratios as defined in Section 4.1.

ARTICLE X: DISPUTE RESOLUTION

10.1 Negotiation: In the event of any dispute arising out of or relating to this Agreement, the Partners shall first attempt to resolve the dispute through good faith negotiation among themselves.
10.2 Mediation: If the dispute cannot be resolved through negotiation within thirty (30) days, the Partners agree to attempt to resolve the dispute through mediation administered by a mutually agreed-upon mediator before resorting to arbitration or litigation.
10.3 Arbitration (Optional): [If the Partners prefer arbitration over court litigation, include:] If mediation is unsuccessful, any controversy or claim arising out of or relating to this Agreement, or the breach thereof, shall be settled by binding arbitration administered in [Specify City, e.g., San Jose Del Monte, Bulacan] in accordance with the arbitration rules of [Specify Body, e.g., the Philippine Dispute Resolution Center, Inc. (PDRCI)], and judgment upon the award rendered by the arbitrator(s) may be entered in any court having jurisdiction thereof.
10.4 Governing Law: This Agreement shall be governed by and construed in accordance with the laws of the Republic of the Philippines.

ARTICLE XI: MISCELLANEOUS

11.1 Notices: Any notice required or permitted under this Agreement shall be in writing and shall be deemed delivered when personally delivered, sent by reputable overnight courier, or sent by registered or certified mail, return receipt requested, postage prepaid, to the addresses of the Partners set forth herein or to such other address as a Partner may designate in writing. Email notice shall be sufficient if receipt is acknowledged.
11.2 Entire Agreement: This Agreement (including Exhibit A) constitutes the entire agreement among the Partners concerning the subject matter hereof and supersedes all prior agreements, understandings, negotiations, and discussions, whether oral or written.
11.3 Amendments: This Agreement may be amended only by a written instrument signed by all Partners.
11.4 Binding Effect: This Agreement shall be binding upon and inure to the benefit of the Partners and their respective heirs, executors, administrators, successors, and permitted assigns.
11.5 Severability: If any provision of this Agreement is held to be invalid, illegal, or unenforceable, the validity, legality, and enforceability of the remaining provisions shall not in any way be affected or impaired thereby.
11.6 Counterparts: This Agreement may be executed in one or more counterparts, each of which shall be deemed an original, but all of which together shall constitute one and the same instrument.

IN WITNESS WHEREOF, the Partners have executed this General Partnership Agreement as of the Effective Date first written above.


(Signature Blocks Placeholder - Add manually or modify script for lines)

_________________________        _________________________
Stella Marie B. Abanono           Nathaniel G. Bingco
(Printed Name)                       (Printed Name)

_________________________        _________________________
Samantha Jane S. Darullo      Louize Christopher S. Galang
(Printed Name)                       (Printed Name)

_________________________        _________________________
Yuri B. Gatbunton                 Neil Vincent E. Jaen
(Printed Name)                       (Printed Name)

_________________________
Myke Jonard S. Sentillas
(Printed Name)


EXHIBIT A

INITIAL CAPITAL CONTRIBUTIONS

As of the Effective Date of the General Partnership Agreement of Drip-o-matic:

| Partner Name                 | Description of Contribution             | Agreed Value (PHP) |
| :--------------------------- | :-------------------------------------- | :----------------- |
| Stella Marie B. Abanono      | [e.g., Cash, Equipment, Marketing Plan] | [PHP Amount]       |
| Nathaniel G. Bingco        | [e.g., Cash, HR Setup Services]         | [PHP Amount]       |
| Samantha Jane S. Darullo   | [e.g., Cash, Financial Model]           | [PHP Amount]       |
| Louize Christopher S. Galang | [e.g., Cash, Prototype Components, Code] | [PHP Amount]       |
| Yuri B. Gatbunton            | [e.g., Cash, Business Plan Concept]     | [PHP Amount]       |
| Neil Vincent E. Jaen         | [e.g., Cash, Operational Workflow]      | [PHP Amount]       |
| Myke Jonard S. Sentillas     | [e.g., Cash, Sales Strategy]            | [PHP Amount]       |
| TOTAL                    |                                         | [PHP Total]    |

*Note: This Exhibit needs to be filled out based on what the partners actually agree to contribute.*
"""

# --- Document Settings ---
OUTPUT_FILENAME = "Drip-o-matic_Partnership_Agreement.docx"
FONT_NAME = 'Arial' # Or 'Times New Roman', etc.
FONT_SIZE_TITLE = 16
FONT_SIZE_SUBTITLE = 14
FONT_SIZE_HEADING = 12
FONT_SIZE_BODY = 10
INDENT_LIST = Inches(0.25)
INDENT_SUB_LIST = Inches(0.5)

# --- Helper function to add Exhibit A Table ---
def add_exhibit_a_table(document, table_lines):
    """Adds the Exhibit A table to the document."""
    # Filter out header separator line and extract data
    header_line = None
    data_rows = []
    for line in table_lines:
        if line.startswith("| :---"): # Skip separator
            continue
        elif line.startswith("| Partner"): # Identify header
             header_line = line
        elif line.startswith("|"): # Data row
            match = re.match(r"\|\s*(.*?)\s*\|\s*(.*?)\s*\|\s*(.*?)\s*\|", line)
            if match:
                data_rows.append(match.groups())

    if not header_line or not data_rows:
        print("Warning: Could not parse Exhibit A table data properly.")
        # Add placeholder if parsing fails
        p = document.add_paragraph("[Exhibit A Table Data - check formatting in source text]")
        return

    # Extract header text
    header_match = re.match(r"\|\s*(.*?)\s*\|\s*(.*?)\s*\|\s*(.*?)\s*\|", header_line)
    if not header_match:
         print("Warning: Could not parse Exhibit A table header.")
         p = document.add_paragraph("[Exhibit A Table Header - check formatting in source text]")
         return
    headers = [h.strip() for h in header_match.groups()]


    # Create table (rows = header + data rows)
    num_rows = 1 + len(data_rows)
    num_cols = len(headers)
    table = document.add_table(rows=num_rows, cols=num_cols)
    table.style = 'Table Grid' # Apply basic grid style
    table.alignment = WD_TABLE_ALIGNMENT.CENTER # Center table on page

    # --- Populate Header Row ---
    hdr_cells = table.rows[0].cells
    for i, header_text in enumerate(headers):
        cell_paragraph = hdr_cells[i].paragraphs[0]
        run = cell_paragraph.add_run(header_text)
        run.bold = True
        run.font.name = FONT_NAME
        run.font.size = Pt(FONT_SIZE_BODY)
        cell_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # --- Populate Data Rows ---
    for row_idx, row_data in enumerate(data_rows, start=1):
        row_cells = table.rows[row_idx].cells
        is_total_row = "TOTAL" in row_data[0] # Check if it's the TOTAL row

        for col_idx, cell_text in enumerate(row_data):
            cell_paragraph = row_cells[col_idx].paragraphs[0]
            run = cell_paragraph.add_run(cell_text.strip())
            run.font.name = FONT_NAME
            run.font.size = Pt(FONT_SIZE_BODY - 1) # Slightly smaller for data
            if is_total_row:
                 run.bold = True # Bold the TOTAL row

            # Align value column right
            if col_idx == num_cols - 1: # Last column (value)
                cell_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            else:
                cell_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # --- Set Column Widths (Optional, adjust as needed) ---
    # Widths are approximate, Word adjusts layout
    widths = (Inches(1.8), Inches(3.0), Inches(1.2))
    try:
        for i, width in enumerate(widths):
            if i < len(table.columns):
                for cell in table.columns[i].cells:
                     cell.width = width # Applies width to all cells in the column
    except IndexError:
         print("Warning: Column width setting issue - check table structure.")

    # Add space after table
    document.add_paragraph()


# --- Main Document Creation ---
document = Document()

# Set default font for the document (optional, styles can override)
style = document.styles['Normal']
font = style.font
font.name = FONT_NAME
font.size = Pt(FONT_SIZE_BODY)

# --- Add Content ---

# Title
p_title = document.add_paragraph()
p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run_title = p_title.add_run("GENERAL PARTNERSHIP AGREEMENT")
run_title.bold = True
run_title.font.size = Pt(FONT_SIZE_TITLE)

# Subtitle
p_subtitle1 = document.add_paragraph()
p_subtitle1.alignment = WD_ALIGN_PARAGRAPH.CENTER
run_subtitle1 = p_subtitle1.add_run("OF")
run_subtitle1.bold = True
run_subtitle1.font.size = Pt(FONT_SIZE_SUBTITLE)

p_subtitle2 = document.add_paragraph()
p_subtitle2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run_subtitle2 = p_subtitle2.add_run("DRIP-O-MATIC")
run_subtitle2.bold = True
run_subtitle2.font.size = Pt(FONT_SIZE_SUBTITLE)
p_subtitle2.paragraph_format.space_after = Pt(18) # Add space after subtitle

# Body Text Processing
lines = agreement_text.strip().split('\n')

# Find start index (after title/subtitle we already added)
start_index = 0
for i, line in enumerate(lines):
    if "Effective Date" in line:
        start_index = i
        break

in_exhibit_a_table = False
exhibit_a_lines = []

# Process the rest of the lines
for line in lines[start_index:]:
    line = line.strip()
    if not line:
        if not in_exhibit_a_table:
            document.add_paragraph() # Add space for empty lines
        continue

    # Detect start and end of Exhibit A Table
    if line.startswith("EXHIBIT A"):
        # Add the heading first
        p = document.add_paragraph()
        p.paragraph_format.space_before = Pt(12)
        run = p.add_run(line)
        run.bold = True
        run.font.size = Pt(FONT_SIZE_HEADING)
        p.paragraph_format.space_after = Pt(6)
        # Then start collecting table lines
        in_exhibit_a_table = True
        exhibit_a_lines = []
        continue # Skip normal processing for this line
    elif in_exhibit_a_table and line.startswith("|"):
        exhibit_a_lines.append(line)
        continue
    elif in_exhibit_a_table and not line.startswith("|"):
        # End of table detected
        add_exhibit_a_table(document, exhibit_a_lines)
        in_exhibit_a_table = False
        # Fall through to process the current line (likely the Note)

    # Process non-table lines
    if not in_exhibit_a_table:
        # Heuristics for formatting
        is_heading = (
            line.isupper() and len(line.split()) < 5 and not line.startswith(('(', '|', '*', '[', '_')) or
            line.startswith("ARTICLE ") or
            line.startswith("RECITALS") or
            line.startswith("IN WITNESS WHEREOF")
        )
        is_subheading = line.endswith(':') and (line[0].isdigit() and '.' in line[:4])
        is_list_item = line.startswith(('1. ', '2. ', '3. ', '4. ', '5. ', '6. ', '7. ', '    *'))
        is_sub_list_item = line.startswith(('(a)', '(b)', '(c)', '(d)', '(e)', '(f)', '(i)', '(ii)', '(iii)', '(iv)', '(v)', '(vi)', '(vii)', '(viii)'))
        is_signature_line = line.startswith('___')
        is_signature_name = line.startswith('(Printed Name)') or line.startswith('(') and ')' in line and len(line) < 30 # Simple check
        is_note = line.startswith('*Note:')

        p = document.add_paragraph()

        if is_heading:
            p.paragraph_format.space_before = Pt(12)
            run = p.add_run(line)
            run.bold = True
            run.font.size = Pt(FONT_SIZE_HEADING)
            p.paragraph_format.space_after = Pt(6)
        elif is_subheading:
            run = p.add_run(line)
            run.bold = True
            # Keep body font size
        elif is_list_item:
            p.paragraph_format.left_indent = INDENT_LIST
            # Add tab stop for better alignment after number (optional)
            # p.paragraph_format.tab_stops.add_tab_stop(Inches(0.5))
            # text = line.split(' ', 1)
            # p.add_run(text[0] + '\t' + text[1]) # Requires adjusting split logic
            p.add_run(line) # Simpler: just indent
        elif is_sub_list_item:
             p.paragraph_format.left_indent = INDENT_SUB_LIST
             p.add_run(line)
        elif is_signature_line:
             p.paragraph_format.space_before = Pt(6)
             p.add_run(line)
        elif is_signature_name:
             # Indent signature names slightly if needed
             p.paragraph_format.left_indent = Inches(0.5) # Adjust as desired
             run = p.add_run(line)
             run.font.size = Pt(FONT_SIZE_BODY -1)
        elif is_note:
             run = p.add_run(line)
             run.italic = True
             run.font.size = Pt(FONT_SIZE_BODY -1)
        else:
            # Regular paragraph
            p.add_run(line)

# --- Final Check for Table at end ---
if in_exhibit_a_table and exhibit_a_lines:
     add_exhibit_a_table(document, exhibit_a_lines)

# --- Save the Document ---
try:
    document.save(OUTPUT_FILENAME)
    print(f"Successfully created '{OUTPUT_FILENAME}'")
except Exception as e:
    print(f"Error creating Word document: {e}")