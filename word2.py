# -*- coding: utf-8 -*-
import os
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

def add_bold_paragraph(document, text_bold, text_normal=""):
    """Adds a paragraph starting with bold text."""
    p = document.add_paragraph()
    p.add_run(text_bold).bold = True
    if text_normal:
        p.add_run(text_normal)
    return p

# --- Create Document ---
document = Document()

# --- Title ---
p_title1 = document.add_paragraph("GENERAL PARTNERSHIP AGREEMENT")
p_title1.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_title1.runs[0].bold = True
p_title1.runs[0].font.size = Pt(14)

p_title2 = document.add_paragraph("OF")
p_title2.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_title2.runs[0].bold = True
p_title2.runs[0].font.size = Pt(14)


p_title3 = document.add_paragraph("DRIP-O-MATIC")
p_title3.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_title3.runs[0].bold = True
p_title3.runs[0].font.size = Pt(14)

document.add_paragraph() # Add some space

# --- Intro ---
document.add_paragraph(
    "This General Partnership Agreement (the \"Agreement\") is made and entered into this 10th day of "
    "January, 2025 (the \"Effective Date\"), by and among the following individuals "
    "(collectively referred to as the \"Partners,\" and individually as a \"Partner\"):"
)
document.add_paragraph() # Add space

# --- Partners List ---
partners = [
    "Stella Marie B. Abanono",
    "Nathaniel G. Bingco",
    "Samantha Jane S. Darullo",
    "Louize Christopher S. Galang",
    "Yuri B. Gatbunton",
    "Neil Vincent E. Jaen",
    "Myke Jonard S. Sentillas"
]
for i, partner in enumerate(partners, 1):
    document.add_paragraph(f"{i}. {partner}", style='List Number')

document.add_paragraph() # Add space

# --- Recitals ---
document.add_heading("RECITALS", level=1)
document.add_paragraph(
    "WHEREAS, the Partners desire to associate themselves as partners in a general partnership for the purpose of "
    "engaging in the business described herein; and"
)
document.add_paragraph(
    "WHEREAS, the Partners wish to establish their respective rights, duties, and liabilities relating to the "
    "Partnership and its business;"
)
document.add_paragraph(
    "NOW, THEREFORE, in consideration of the mutual covenants contained herein, the Partners agree as follows:"
)
document.add_paragraph() # Add space

# --- ARTICLE I: FORMATION, NAME, AND PURPOSE ---
document.add_heading("ARTICLE I: FORMATION, NAME, AND PURPOSE", level=1)
add_bold_paragraph(document, "1.1 Formation: ", "The Partners hereby form a general partnership (the \"Partnership\") pursuant to the laws of the Republic of the Philippines.")
add_bold_paragraph(document, "1.2 Name: ", "The name of the Partnership shall be Drip-o-matic. The business of the Partnership shall be conducted under this name or such other names as the Partners may unanimously agree upon.")
add_bold_paragraph(document, "1.3 Purpose: ", "The primary purpose of the Partnership is to design, develop, manufacture, market, and sell the \"Drip-o-matic\" automatic plant care system and related products and services, including automatic irrigation, soil moisture monitoring, temperature regulation, plant misting, and solar power integration, aiming to provide convenient and efficient plant care solutions. The Partnership may also engage in any other lawful business activities related to its primary purpose as the Partners may unanimously agree.")
add_bold_paragraph(document, "1.4 Principal Place of Business: ", "The principal place of business of the Partnership shall initially be located at STI Academic Center Altaraza Town Center, Tungkong Mangga, CSJDM, or such other location(s) as the Partners may unanimously determine.")
document.add_paragraph()

# --- ARTICLE II: TERM ---
document.add_heading("ARTICLE II: TERM", level=1)
add_bold_paragraph(document, "2.1 Duration: ", "The Partnership shall commence on the Effective Date and shall continue until dissolved and terminated in accordance with the provisions of this Agreement or by operation of law.")
document.add_paragraph()

# --- ARTICLE III: CAPITAL CONTRIBUTIONS ---
document.add_heading("ARTICLE III: CAPITAL CONTRIBUTIONS", level=1)
add_bold_paragraph(document, "3.1 Initial Contributions: ", "Each Partner shall contribute capital to the Partnership as mutually agreed upon. The description and agreed value of the initial capital contributions (which may include cash, property, services, or expertise) made by each Partner shall be detailed in Exhibit A, attached hereto and incorporated herein by reference.")
add_bold_paragraph(document, "3.2 Additional Contributions: ", "No Partner shall be required to make additional capital contributions unless unanimously agreed upon in writing by all Partners.")
add_bold_paragraph(document, "3.3 No Interest: ", "No Partner shall be entitled to receive interest on their capital contributions.")
add_bold_paragraph(document, "3.4 Capital Accounts: ", "An individual capital account shall be maintained for each Partner, reflecting their contributions, share of profits and losses, and distributions.")
document.add_paragraph()

# --- ARTICLE IV: PROFITS, LOSSES, AND DISTRIBUTIONS ---
document.add_heading("ARTICLE IV: PROFITS, LOSSES, AND DISTRIBUTIONS", level=1)
add_bold_paragraph(document, "4.1 Allocation of Profits and Losses: ", "Unless otherwise unanimously agreed in writing, the net profits and losses of the Partnership (after deduction of all expenses, including any salaries approved under Article VI) shall be allocated among the Partners in equal shares (i.e., 1/7th each).")
add_bold_paragraph(document, "4.2 Distributions: ", "Distributions of cash or other property shall be made to the Partners at such times and in such amounts as the Partners may unanimously determine, typically reviewed on a quarterly basis, provided that such distributions do not impair the ability of the Partnership to meet its obligations. Distributions shall be made in proportion to the Partners' respective shares of profits as defined in Section 4.1.")
document.add_paragraph()

# --- ARTICLE V: MANAGEMENT AND DUTIES ---
document.add_heading("ARTICLE V: MANAGEMENT AND DUTIES", level=1)
add_bold_paragraph(document, "5.1 Management: ", "All Partners shall have the right to participate in the management and control of the Partnership's business.")
add_bold_paragraph(document, "5.2 Officer Roles: ", "The Partners agree to designate initial officer roles to manage specific operational areas, subject to the overall authority of the Partnership:")
officers = [
    "Chief Executive Officer (CEO): Yuri B. Gatbunton",
    "Chief Operating Officer (COO): Neil Vincent E. Jaen",
    "Chief Financial Officer (CFO): Samantha Jane S. Darullo",
    "Chief Marketing Officer (CMO): Stella Marie B. Abanono",
    "Chief Technology Officer (CTO): Louize Christopher S. Galang",
    "Chief Sales Officer (CSO): Myke Jonard S. Sentillas",
    "Chief Human Resources Officer (CHRO): Nathaniel G. Bingco"
]
for officer in officers:
    document.add_paragraph(officer, style='List Bullet')

document.add_paragraph(
    "The specific duties associated with these roles shall be as generally understood for such positions and as further "
    "defined by unanimous agreement of the Partners. These are internal management roles; all Partners remain general "
    "partners with associated rights and liabilities."
)

add_bold_paragraph(document, "5.3 Decision Making:")
p53a = document.add_paragraph()
p53a.add_run("(a) Ordinary Course Decisions: ").bold = True
p53a.add_run("Decisions arising in the ordinary course of the Partnership's business may be made by the designated officer responsible for that area, subject to review by the Partners.")

p53b = document.add_paragraph()
p53b.add_run("(b) Major Decisions: ").bold = True
p53b.add_run("The following decisions shall require the unanimous written consent of all Partners:")

major_decisions = [
    "Admitting a new Partner;",
    "Amending this Agreement;",
    "Selling, leasing, or disposing of all or substantially all of the Partnership's assets;",
    "Incurring debt or making expenditures outside the approved budget or above a threshold amount of PHP 50,000;", # Threshold filled
    "Changing the primary purpose or nature of the Partnership's business;",
    "Merging or consolidating the Partnership with another entity;",
    "Filing for bankruptcy or initiating dissolution;",
    "Approving annual budgets and major strategic plans."
]
for i, decision in enumerate(major_decisions):
    document.add_paragraph(f"      ({chr(105+i)}) {decision}") # Using chr(105) for 'i'

add_bold_paragraph(document, "5.4 Devotion of Time: ", "Each Partner shall devote such time and attention to the Partnership business as is reasonably necessary for the discharge of their agreed-upon duties and responsibilities.")
document.add_paragraph()

# --- ARTICLE VI: SALARIES AND COMPENSATION ---
document.add_heading("ARTICLE VI: SALARIES AND COMPENSATION", level=1)
add_bold_paragraph(document, "6.1 Salaries: ", "Initially, no Partner shall receive a fixed salary for services rendered to the Partnership. Partners shall be compensated through their share of profits as provided in Article IV. The Partners may unanimously agree in writing to provide salaries or guaranteed payments to one or more Partners at a future date.")
add_bold_paragraph(document, "6.2 Expense Reimbursement: ", "Partners shall be reimbursed for reasonable and necessary out-of-pocket expenses incurred directly on behalf of the Partnership business, provided such expenses are properly documented and approved according to procedures established by the Partners.")
document.add_paragraph()

# --- ARTICLE VII: BOOKS, RECORDS, AND BANKING ---
document.add_heading("ARTICLE VII: BOOKS, RECORDS, AND BANKING", level=1)
add_bold_paragraph(document, "7.1 Books and Records: ", "The Partnership shall maintain complete and accurate books and records of its business and financial affairs at its principal place of business, in accordance with generally accepted accounting principles (GAAP). Each Partner shall have the right to inspect and copy these records at reasonable times.")
add_bold_paragraph(document, "7.2 Fiscal Year: ", "The fiscal year of the Partnership shall end on December 31st.")
# **MODIFIED 7.3 HERE**
add_bold_paragraph(document, "7.3 Banking: ", "All funds of the Partnership shall be deposited in the Partnership's name in such bank account(s) as selected by the Partners. Withdrawals from such accounts shall be made upon the joint signatures of the designated Chief Executive Officer (CEO) and Chief Financial Officer (CFO).")
document.add_paragraph()

# --- ARTICLE VIII: PARTNER WITHDRAWAL, RETIREMENT, DEATH ---
document.add_heading("ARTICLE VIII: PARTNER WITHDRAWAL, RETIREMENT, DEATH", level=1)
add_bold_paragraph(document, "8.1 Voluntary Withdrawal: ", "Any Partner may withdraw from the Partnership upon providing at least ninety (90) days' prior written notice to the other Partners.")
add_bold_paragraph(document, "8.2 Buy-Out Option: ", "Upon the withdrawal, retirement, death, or incapacity of a Partner (the \"Departing Partner\"), the remaining Partners shall have the option, exercisable within sixty (60) days of the event, to purchase the Departing Partner's interest in the Partnership.")
add_bold_paragraph(document, "8.3 Valuation: ", "The purchase price for the Departing Partner's interest shall be the value of their capital account as of the date of departure, adjusted for profits/losses up to that date, plus or minus any other adjustments agreed upon by the Departing Partner (or their estate) and the remaining Partners. If an agreement on value cannot be reached, the value shall be determined by an independent appraiser mutually acceptable to both parties.")
add_bold_paragraph(document, "8.4 Payment Terms: ", "The purchase price shall be paid by the remaining Partners (or the Partnership, if agreed) over a period not to exceed twelve (12) months, with terms to be mutually agreed upon.") # Period filled
add_bold_paragraph(document, "8.5 Dissolution upon Departure: ", "If the remaining Partners do not exercise the option to purchase the Departing Partner's interest, the Partnership may be dissolved in accordance with Article IX.")
document.add_paragraph()

# --- ARTICLE IX: DISSOLUTION AND LIQUIDATION ---
document.add_heading("ARTICLE IX: DISSOLUTION AND LIQUIDATION", level=1)
add_bold_paragraph(document, "9.1 Events of Dissolution: ", "The Partnership shall be dissolved upon the occurrence of any of the following:")
dissolution_events = [
    "The unanimous written agreement of all Partners to dissolve;",
    "The expiration of the term, if specified and not extended;",
    "The sale or disposition of all or substantially all Partnership assets;",
    "An event making it unlawful for the business of the Partnership to be carried on;",
    "Entry of a decree of judicial dissolution;",
    "Failure of remaining partners to purchase a Departing Partner's interest as per Section 8.5, if they elect dissolution."
]
for i, event in enumerate(dissolution_events):
     document.add_paragraph(f"      ({chr(97+i)}) {event}") # Using chr(97) for 'a'

add_bold_paragraph(document, "9.2 Liquidation: ", "Upon dissolution, the Partnership shall cease carrying on business except as necessary for winding up its affairs. The Partners (or a designated liquidating Partner/trustee) shall proceed with reasonable promptness to liquidate the Partnership's assets.")
add_bold_paragraph(document, "9.3 Distribution of Assets: ", "The proceeds from liquidation shall be applied in the following order of priority:")
distribution_priority = [
    "To pay debts and liabilities of the Partnership owed to third-party creditors;",
    "To pay debts and liabilities owed to Partners other than for capital and profits;",
    "To repay Partners' capital contributions (adjusted for profits/losses);",
    "Any remaining amount shall be distributed to the Partners according to their profit-sharing ratios as defined in Section 4.1."
]
for i, item in enumerate(distribution_priority):
    document.add_paragraph(f"      ({chr(97+i)}) {item}") # Using chr(97) for 'a'
document.add_paragraph()

# --- ARTICLE X: DISPUTE RESOLUTION ---
document.add_heading("ARTICLE X: DISPUTE RESOLUTION", level=1)
add_bold_paragraph(document, "10.1 Negotiation: ", "In the event of any dispute arising out of or relating to this Agreement, the Partners shall first attempt to resolve the dispute through good faith negotiation among themselves.")
add_bold_paragraph(document, "10.2 Mediation: ", "If the dispute cannot be resolved through negotiation within thirty (30) days, the Partners agree to attempt to resolve the dispute through mediation administered by a mutually agreed-upon mediator before resorting to arbitration or litigation.")
# **MODIFIED 10.3 HEADING AND TEXT**
add_bold_paragraph(document, "10.3 Arbitration: ", "If mediation is unsuccessful, any controversy or claim arising out of or relating to this Agreement, or the breach thereof, shall be settled by binding arbitration administered in San Jose Del Monte, Bulacan in accordance with the arbitration rules of the Philippine Dispute Resolution Center, Inc. (PDRCI), and judgment upon the award rendered by the arbitrator(s) may be entered in any court having jurisdiction thereof.") # Arbitration details filled, "(Optional)" removed from heading
add_bold_paragraph(document, "10.4 Governing Law: ", "This Agreement shall be governed by and construed in accordance with the laws of the Republic of the Philippines.")
document.add_paragraph()

# --- ARTICLE XI: MISCELLANEOUS ---
document.add_heading("ARTICLE XI: MISCELLANEOUS", level=1)
add_bold_paragraph(document, "11.1 Notices: ", "Any notice required or permitted under this Agreement shall be in writing and shall be deemed delivered when personally delivered, sent by reputable overnight courier, or sent by registered or certified mail, return receipt requested, postage prepaid, to the addresses of the Partners set forth herein or to such other address as a Partner may designate in writing. Email notice shall be sufficient if receipt is acknowledged.")
add_bold_paragraph(document, "11.2 Entire Agreement: ", "This Agreement (including Exhibit A) constitutes the entire agreement among the Partners concerning the subject matter hereof and supersedes all prior agreements, understandings, negotiations, and discussions, whether oral or written.")
add_bold_paragraph(document, "11.3 Amendments: ", "This Agreement may be amended only by a written instrument signed by all Partners.")
add_bold_paragraph(document, "11.4 Binding Effect: ", "This Agreement shall be binding upon and inure to the benefit of the Partners and their respective heirs, executors, administrators, successors, and permitted assigns.")
add_bold_paragraph(document, "11.5 Severability: ", "If any provision of this Agreement is held to be invalid, illegal, or unenforceable, the validity, legality, and enforceability of the remaining provisions shall not in any way be affected or impaired thereby.")
add_bold_paragraph(document, "11.6 Counterparts: ", "This Agreement may be executed in one or more counterparts, each of which shall be deemed an original, but all of which together shall constitute one and the same instrument.")
document.add_paragraph()

# --- WITNESS ---
p_witness = document.add_paragraph()
p_witness.add_run("IN WITNESS WHEREOF").bold = True
p_witness.add_run(", the Partners have executed this General Partnership Agreement as of the Effective Date first written above.")
document.add_paragraph()
document.add_paragraph() # Extra space where placeholder was

# --- Signature Blocks ---
sig_line = "_________________________"
num_partners = len(partners)
for i in range(0, num_partners, 2):
    p_sig = document.add_paragraph()
    p_sig.add_run(sig_line)
    p_sig.add_run("\t\t") # Add tabs for spacing
    if i + 1 < num_partners:
        p_sig.add_run(sig_line)

    p_name = document.add_paragraph()
    p_name.add_run(partners[i])
    p_name.add_run("\t\t\t") # Add tabs for spacing
    if i + 1 < num_partners:
         p_name.add_run(partners[i+1])

    p_printed = document.add_paragraph()
    p_printed.add_run("(Printed Name)")
    p_printed.add_run("\t\t\t") # Add tabs for spacing
    if i + 1 < num_partners:
        p_printed.add_run("(Printed Name)")

    document.add_paragraph() # Add space between signature pairs

# Add a page break before Exhibit A
document.add_page_break()

# --- EXHIBIT A ---
document.add_heading("EXHIBIT A", level=1)
document.add_paragraph()
document.add_heading("INITIAL CAPITAL CONTRIBUTIONS", level=2)
document.add_paragraph("As of the Effective Date of the General Partnership Agreement of Drip-o-matic:")
document.add_paragraph()

# --- Exhibit A Table ---
exhibit_data = [
    ("Stella Marie B. Abanono", "Cash: PHP 20,000; Initial Marketing & Branding Plan (Services)", "PHP 50,000"),
    ("Nathaniel G. Bingco", "Cash: PHP 25,000; Draft HR Policies & Setup Plan (Services)", "PHP 50,000"),
    ("Samantha Jane S. Darullo", "Cash: PHP 30,000; Financial Model & Budget Setup (Services)", "PHP 50,000"),
    ("Louize Christopher S. Galang", "Cash: PHP 15,000; Prototype Components & Initial Code (Property/IP)", "PHP 50,000"),
    ("Yuri B. Gatbunton", "Cash: PHP 25,000; Detailed Business Plan & Strategy Doc (Services)", "PHP 50,000"),
    ("Neil Vincent E. Jaen", "Cash: PHP 20,000; Operational Workflow Design (Services)", "PHP 50,000"),
    ("Myke Jonard S. Sentillas", "Cash: PHP 25,000; Initial Sales Strategy & Market Analysis (Services)", "PHP 50,000"),
]

table = document.add_table(rows=1, cols=3)
table.style = 'Table Grid' # Apply grid lines

# Header Row
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Partner Name'
hdr_cells[1].text = 'Description of Contribution'
hdr_cells[2].text = 'Agreed Value (PHP)'
for cell in hdr_cells:
    cell.paragraphs[0].runs[0].bold = True
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

# Data Rows
for name, desc, value in exhibit_data:
    row_cells = table.add_row().cells
    row_cells[0].text = name
    row_cells[1].text = desc
    row_cells[2].text = value
    row_cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

# Total Row
row_cells = table.add_row().cells
row_cells[0].text = 'TOTAL'
row_cells[0].paragraphs[0].runs[0].bold = True
row_cells[1].text = '' # Empty description cell
row_cells[2].text = 'PHP 350,000'
row_cells[2].paragraphs[0].runs[0].bold = True
row_cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

# Set column widths
widths = (Inches(1.75), Inches(3.5), Inches(1.25))
for row in table.rows:
    for idx, width in enumerate(widths):
        try: # Added try-except as sometimes width setting can fail on complex tables
            row.cells[idx].width = width
        except Exception as e:
            print(f"Warning: Could not set width for cell {idx} in a row. Error: {e}")


document.add_paragraph()
p_note = document.add_paragraph()
p_note.add_run("*Note: ").italic = True
p_note.add_run("The non-cash contributions listed above have been valued by unanimous agreement of the Partners as of the Effective Date.").italic = True


# --- Save Document ---
file_name = "Drip-o-matic_Partnership_Agreement_Final_2025_v2.docx" # Renamed slightly
try:
    document.save(file_name)
    print(f"Document '{file_name}' created successfully in script directory: {os.path.abspath(file_name)}")
except Exception as e:
    print(f"Error saving document: {e}")